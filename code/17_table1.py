#!/usr/bin/env python3
"""
17_table1.py
============
Emit manuscript Table 1 as a Word document and a CSV.

Inputs
    output/validation/manuscript_numbers.json

Outputs
    output/tables/Table1_Capacity_by_Rurality_and_EDI.docx
    output/tables/Table1_Capacity_by_Rurality_and_EDI.csv

Every value is read from manuscript_numbers.json, which is the same file the
validation gate checks the manuscript against cell by cell. Nothing is
recomputed here, so the artifact cannot disagree with the paper: if a number
moves, it moves in both places or the gate fails.

Note that this is a different table from the SVI-quartile descriptives written
by 07_publication_outputs.py. This one is the manuscript's Table 1, stratified
by rurality and by EDI quintile.
"""

from __future__ import annotations

import csv
import json
import os

from docx import Document
from docx.shared import Pt

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
NUMBERS = os.path.join(BASE_DIR, "output", "validation", "manuscript_numbers.json")
TABLES = os.path.join(BASE_DIR, "output", "tables")
STEM = "Table1_Capacity_by_Rurality_and_EDI"

TITLE = ("Table 1. Accredited cardiac imaging capacity by rurality and economic "
         "deprivation, United States, 2026")

HEADERS = ["Stratum",
           "Counties, No.",
           "Adults ≥45 y, millions",
           "CMR facilities, No.",
           "Counties with ≥1 CMR, No. (%)",
           "CMR rate, mean per 100,000",
           "CCT facilities, No.",
           "Counties with ≥1 CCT, No. (%)",
           "CCT rate, mean per 100,000"]

#: (json key, printed label). None marks a subheading row with no data.
LAYOUT = [("all_counties", "All counties"),
          (None, "Rurality"),
          ("metropolitan", "Metropolitan (RUCC 1-3)"),
          ("nonmetropolitan", "Nonmetropolitan (RUCC 4-9)"),
          (None, "EDI quintile"),
          ("edi_q1", "Q1 (least deprived)"),
          ("edi_q2", "Q2"),
          ("edi_q3", "Q3"),
          ("edi_q4", "Q4"),
          ("edi_q5", "Q5 (most deprived)")]


def build_rows(t1):
    """Table body as printable strings, in manuscript order."""
    rows = []
    for key, label in LAYOUT:
        if key is None:
            rows.append([label] + [""] * (len(HEADERS) - 1))
            continue
        b = t1[key]
        rows.append([
            label,
            f"{b['counties']:,}",
            f"{b['adults_millions']:.1f}",
            f"{b['cmr_facilities']:,}",
            f"{b['counties_ge1_cmr']:,} ({b['counties_ge1_cmr_pct']:.1f})",
            f"{b['cmr_rate_mean']:.2f}",
            f"{b['cct_facilities']:,}",
            f"{b['counties_ge1_cct']:,} ({b['counties_ge1_cct_pct']:.1f})",
            f"{b['cct_rate_mean']:.2f}",
        ])
    return rows


def footnotes(t1):
    """Provenance the table cannot be read correctly without."""
    excluded = t1["all_counties"]["counties"] - t1["all_counties"]["rate_eligible"]
    return [
        "CMR = cardiac magnetic resonance; CCT = cardiac computed tomography; "
        "RUCC = Rural-Urban Continuum Code (USDA ERS 2023); EDI = Economic "
        "Deprivation Index.",
        "Metropolitan is RUCC 1-3, nonmetropolitan RUCC 4-9.",
        "Rate = accredited facilities per 100,000 adults aged 45 years and older.",
        # The counts and the rates use different denominators on purpose, and a
        # reader who assumes otherwise will not be able to reproduce the rates.
        f"Facility and county counts include every county in the stratum. Mean "
        f"rates are calculated over the rate-eligible counties only, because a "
        f"per-head rate is not defined for a county with fewer than 1,000 adults "
        f"aged 45 and older ({excluded:,} counties excluded; "
        f"{t1['all_counties']['rate_eligible']:,} of "
        f"{t1['all_counties']['counties']:,} retained).",
        "EDI quintiles are formed among rate-eligible counties with an EDI value, "
        "so the rurality strata and the quintile strata do not sum to the same "
        "total.",
        "Population denominators are American Community Survey 5-year estimates, "
        "2019-2023.",
    ]


def write_csv(rows, notes, path):
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(HEADERS)
        w.writerows(rows)
        w.writerow([])
        for n in notes:
            w.writerow([n])


def write_docx(rows, notes, path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9)

    doc.add_heading(TITLE, level=1)

    t = doc.add_table(rows=1, cols=len(HEADERS))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(HEADERS):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for row in rows:
        cells = t.add_row().cells
        is_subhead = all(v == "" for v in row[1:])
        for i, v in enumerate(row):
            cells[i].text = v
            if is_subhead:
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.italic = True

    doc.add_paragraph()
    for n in notes:
        p = doc.add_paragraph(n)
        p.runs[0].font.size = Pt(8)

    doc.save(path)


def main():
    print("=" * 72)
    print("  MANUSCRIPT TABLE 1")
    print("=" * 72)

    if not os.path.exists(NUMBERS):
        raise SystemExit("manuscript_numbers.json not found. "
                         "Run 12_manuscript_numbers.py first.")

    with open(NUMBERS, encoding="utf-8") as f:
        t1 = json.load(f)["table1"]

    rows = build_rows(t1)
    notes = footnotes(t1)

    os.makedirs(TABLES, exist_ok=True)
    write_csv(rows, notes, os.path.join(TABLES, STEM + ".csv"))
    write_docx(rows, notes, os.path.join(TABLES, STEM + ".docx"))

    width = max(len(r[0]) for r in rows)
    for r in rows:
        print(f"  {r[0]:<{width}}  " + "  ".join(f"{v:>10}" for v in r[1:]))

    print(f"\n  wrote {STEM}.docx and {STEM}.csv")


if __name__ == "__main__":
    main()
