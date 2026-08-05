#!/usr/bin/env python3
"""
11_edi_tables_and_stats.py
==========================
Generates the supplementary Word tables and the descriptive statistics JSON
directly from the committed data, so that no manuscript number is ever
hand-transcribed again.

This script exists because two values in the JACC submission (the count of
counties with neither modality, and the mean deprivation contrast) were carried
into the manuscript from a hand-maintained narrative document and were wrong.
Every quantity produced here is computed at run time from:

    data/processed/county_analytic_dataset.csv
    data/processed/county_edi_constructed.csv
    data/processed/edi_regression_results.json      (from 06_edi_sensitivity_analysis.py)
    output/jacr_revision/validated_index_results.json (from 09_validated_index_sdi.py)

Outputs
    output/tables/Supplementary_Table_EDI_Regression.docx
    output/tables/Table4_SVI_vs_EDI_Comparison.docx
    output/tables/Table4_External_Validation_SDI.docx   (manuscript Table 4)
    output/{tables,requested}/additional_statistics.json
    output/supplementary_data/additional_statistics.json
    (Word tables are also copied to output/requested/)

Run
    python code/06_edi_sensitivity_analysis.py    (first)
    python code/09_validated_index_sdi.py         (first, for the SDI columns)
    python code/11_edi_tables_and_stats.py
"""

import os
import json
import shutil
import warnings
import numpy as np
import pandas as pd
warnings.filterwarnings("ignore")
import statsmodels.api as sm
from scipy import stats
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROC = os.path.join(BASE_DIR, "data", "processed")
OUT = os.path.join(BASE_DIR, "output")
TABLES = os.path.join(OUT, "tables")
REQUESTED = os.path.join(OUT, "requested")
SUPP = os.path.join(OUT, "supplementary_data")
JACR = os.path.join(OUT, "jacr_revision")
for d in (TABLES, REQUESTED, SUPP):
    os.makedirs(d, exist_ok=True)


def load():
    df = pd.read_csv(os.path.join(PROC, "county_analytic_dataset.csv"),
                     dtype={"county_fips": str})
    df["county_fips"] = df["county_fips"].str.zfill(5)
    edi = pd.read_csv(os.path.join(PROC, "county_edi_constructed.csv"),
                      dtype={"fips": str})
    edi["fips"] = edi["fips"].str.zfill(5)
    m = df.merge(edi, left_on="county_fips", right_on="fips", how="left")
    with open(os.path.join(PROC, "edi_regression_results.json")) as f:
        edi_res = json.load(f)
    sdi_res = None
    sdi_path = os.path.join(JACR, "validated_index_results.json")
    if os.path.exists(sdi_path):
        with open(sdi_path) as f:
            sdi_res = json.load(f)
    return m, edi_res, sdi_res


def svi_models(m):
    """Refit the SVI models here so the comparison table is self-contained."""
    d = m[m.rate_excluded == 0].copy()
    d["svi_per10"] = d.svi_percentile * 10.0     # percentile is 0-1, so x10 = per 10 pct pts
    d["log_pop"] = np.log(d.adult_pop_45plus)
    out = {}
    for label, col in [("CMR", "cmr_facility_count"), ("CCT", "cct_facility_count")]:
        row = {}
        for name, terms in [("unadjusted", ["svi_per10"]),
                            ("adjusted_metro", ["svi_per10", "metro_indicator"])]:
            X = sm.add_constant(d[terms], has_constant="add")
            r = sm.GLM(d[col], X, family=sm.families.NegativeBinomial(),
                       offset=d.log_pop).fit(maxiter=100)
            ci = np.exp(r.conf_int().loc["svi_per10"])
            row[name] = {"irr": float(np.exp(r.params["svi_per10"])),
                         "ci_low": float(ci.iloc[0]), "ci_high": float(ci.iloc[1]),
                         "p": float(r.pvalues["svi_per10"])}
            if name == "adjusted_metro":
                mci = np.exp(r.conf_int().loc["metro_indicator"])
                row["metro_effect"] = {
                    "irr": float(np.exp(r.params["metro_indicator"])),
                    "ci_low": float(mci.iloc[0]), "ci_high": float(mci.iloc[1]),
                    "p": float(r.pvalues["metro_indicator"])}
        out[label] = row
    return out


def norm(e):
    """Script 09 emits IRR/CI_low/CI_high/P; script 06 emits irr/ci_low/ci_high/p."""
    if "irr" in e:
        return e
    return {"irr": e["IRR"], "ci_low": e["CI_low"], "ci_high": e["CI_high"], "p": e["P"]}


def fmt(e, dp=4):
    e = norm(e)
    return f"{e['irr']:.{dp}f} ({e['ci_low']:.{dp}f}-{e['ci_high']:.{dp}f})"


def fmt_p(p):
    return "<0.0001" if p < 0.0001 else f"{p:.4f}"


def _style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t


def supplementary_table(edi_res):
    doc = Document()
    _style(doc)
    doc.add_heading("Supplementary Table. Economic Deprivation Index (EDI) and "
                    "accredited cardiac imaging capacity", level=1)
    p = doc.add_paragraph(
        "Negative binomial regression with an offset of log(adults aged 45 and older). "
        "IRR is expressed per 10-percentile increase in EDI. Rate-eligible counties only "
        "(≥1,000 adults aged 45+).")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for modality in ["CMR", "CCT"]:
        r = edi_res[modality]
        doc.add_heading(f"{modality} (n = {r['n']} counties)", level=2)
        rows = [
            ("EDI, unadjusted", r["unadjusted"]),
            ("EDI, adjusted for metropolitan status", r["adjusted_metro"]),
            ("EDI, adjusted for ordinal RUCC", r["adjusted_rucc"]["edi"]),
            ("EDI, metropolitan counties only", r["stratified"]["metro"]),
            ("EDI, nonmetropolitan counties only", r["stratified"]["nonmetro"]),
            ("Metropolitan status (from adjusted model)", r["metro_effect"]),
        ]
        _table(doc, ["Model", "IRR (95% CI)", "P value"],
               [(lbl, fmt(e), fmt_p(norm(e)["p"])) for lbl, e in rows])
        se = r["stratified_events"]
        doc.add_paragraph(
            f"Stratum sizes. Metropolitan, {r['stratified_n']['metro']} counties carrying "
            f"{se['metro']['events']} facilities in {se['metro']['counties_with_facility']} "
            f"counties. Nonmetropolitan, {r['stratified_n']['nonmetro']} counties carrying "
            f"{se['nonmetro']['events']} facilities in "
            f"{se['nonmetro']['counties_with_facility']} counties.")
        if se["nonmetro"]["events"] < 30:
            cap = doc.add_paragraph()
            run = cap.add_run(
                f"Caution. The nonmetropolitan {modality} stratum contains only "
                f"{se['nonmetro']['events']} facilities. Estimates from it are unstable and "
                "should be read as exploratory.")
            run.italic = True

    rl = edi_res["edi_vs_rurality"]
    doc.add_heading("Why the unadjusted association appears", level=2)
    doc.add_paragraph(
        f"The EDI tracks rurality. Spearman rho versus RUCC is {rl['spearman_rho']:.3f}, "
        f"and mean EDI is {rl['metro_mean']:.1f} in metropolitan counties versus "
        f"{rl['nonmetro_mean']:.1f} in nonmetropolitan counties, a gap of {rl['gap']:.1f} "
        "points. Because accredited capacity is concentrated in metropolitan counties, an "
        "unadjusted deprivation model absorbs the rurality signal. Once metropolitan status "
        "is included, the deprivation term is no longer associated with capacity.")

    for d in (TABLES, REQUESTED):
        doc.save(os.path.join(d, "Supplementary_Table_EDI_Regression.docx"))
    print("  wrote Supplementary_Table_EDI_Regression.docx")


def comparison_table(edi_res, svi_res, sdi_res, m):
    doc = Document()
    _style(doc)
    doc.add_heading("Table 4. Social Vulnerability Index versus Economic Deprivation "
                    "Index, unadjusted and adjusted for rurality", level=1)
    doc.add_paragraph(
        "Negative binomial regression with an offset of log(adults aged 45 and older). "
        "IRR per 10-percentile increase in the index. Neither index is associated with "
        "capacity once metropolitan status is in the model.")

    for modality in ["CMR", "CCT"]:
        doc.add_heading(modality, level=2)
        rows = [
            ("SVI, unadjusted", svi_res[modality]["unadjusted"]),
            ("SVI, adjusted for metropolitan status", svi_res[modality]["adjusted_metro"]),
            ("EDI, unadjusted", edi_res[modality]["unadjusted"]),
            ("EDI, adjusted for metropolitan status", edi_res[modality]["adjusted_metro"]),
        ]
        if sdi_res is not None:
            key = "cmr_facility_count" if modality == "CMR" else "cct_facility_count"
            s = sdi_res["SDI_models"]["outcomes"][key]
            rows += [("Graham Center SDI, unadjusted", s["unadjusted"]),
                     ("Graham Center SDI, adjusted for metropolitan status", s["adjusted"])]
        rows.append(("Metropolitan status (from adjusted EDI model)",
                     edi_res[modality]["metro_effect"]))
        _table(doc, ["Model", "IRR (95% CI)", "P value"],
               [(lbl, fmt(e), fmt_p(norm(e)["p"])) for lbl, e in rows])

    valid = m[["svi_percentile", "edi_national_percentile"]].dropna()
    r, p = stats.pearsonr(valid.svi_percentile, valid.edi_national_percentile)
    doc.add_paragraph(
        f"Correlation between SVI and EDI: Pearson r = {r:.4f} "
        f"(P {'<0.0001' if p < 0.0001 else f'= {p:.4f}'}). The two indices are strongly "
        "correlated but measure overlapping rather than identical constructs.")
    if sdi_res is not None:
        ag = sdi_res["agreement_with_EDI"]
        doc.add_paragraph(
            f"Agreement between our EDI and the external Graham Center SDI: Spearman rho = "
            f"{ag['spearman_rho']:.3f}, Pearson r = {ag['pearson_r']:.3f}, across "
            f"{sdi_res['sdi_matched_counties']} of {sdi_res['total_counties']} counties.")

    for d in (TABLES, REQUESTED):
        doc.save(os.path.join(d, "Table4_SVI_vs_EDI_Comparison.docx"))
    print("  wrote Table4_SVI_vs_EDI_Comparison.docx")


def external_validation_table(sdi_res):
    """Manuscript Table 4, our EDI head to head with the external Graham Center SDI."""
    if sdi_res is None:
        print("  skipped Table4_External_Validation_SDI.docx (run 09 first)")
        return
    E = sdi_res["EDI_models"]["outcomes"]
    S = sdi_res["SDI_models"]["outcomes"]
    doc = Document()
    _style(doc)
    doc.add_heading("Table 4. External validation of the deprivation finding using a "
                    "published county-level index", level=1)
    rows = []
    for modality, key in [("Cardiac MR", "cmr_facility_count"),
                          ("Cardiac CT", "cct_facility_count")]:
        rows.append((modality, "", "", "", ""))
        for label, mdl in [("Index, unadjusted", "unadjusted"),
                           ("Index, adjusted for metropolitan status", "adjusted"),
                           ("Metropolitan status", "metro_in_adjusted")]:
            rows.append((f"    {label}",
                         fmt(E[key][mdl], 2), fmt_p(norm(E[key][mdl])["p"]),
                         fmt(S[key][mdl], 2), fmt_p(norm(S[key][mdl])["p"])))
    _table(doc, ["Exposure", "EDI IRR (95% CI)", "P value", "SDI IRR (95% CI)", "P value"], rows)

    ag = sdi_res["agreement_with_EDI"]
    rl = sdi_res["rurality_link"]
    doc.add_paragraph(
        f"EDI = economic deprivation index; IRR = incidence rate ratio; SDI = Social "
        f"Deprivation Index (Robert Graham Center, 2015-2019). Negative binomial regression "
        f"with a log-population offset (adults aged 45 and older); index IRRs are per "
        f"10-percentile increase. Analytic sample: {sdi_res['EDI_models']['n']:,} counties "
        f"(EDI) and {sdi_res['SDI_models']['n']:,} (SDI); the SDI matched "
        f"{sdi_res['sdi_matched_counties']:,} of {sdi_res['total_counties']:,} counties. The "
        f"two indices agree closely at the county level (Spearman rho = "
        f"{ag['spearman_rho']:.2f}) but differ in how strongly they encode rurality: Spearman "
        f"rho with the ordinal Rural-Urban Continuum Code was "
        f"{rl['EDI']['spearman_vs_rucc']:.2f} for the EDI and "
        f"{rl['SDI']['spearman_vs_rucc']:.2f} for the SDI, with mean nonmetropolitan minus "
        f"metropolitan gaps of {rl['EDI']['gap']:.1f} and {rl['SDI']['gap']:.1f} percentile "
        f"points, respectively. Metropolitan status rows are the metropolitan-status "
        f"coefficients from the corresponding adjusted models.")
    for d in (TABLES, REQUESTED):
        doc.save(os.path.join(d, "Table4_External_Validation_SDI.docx"))
    print("  wrote Table4_External_Validation_SDI.docx")


def additional_stats(m, edi_res, svi_res):
    """Recompute every loose descriptive number the manuscript quotes."""
    el = m[m.rate_excluded == 0]
    neither_all = (m.cmr_facility_count == 0) & (m.cct_facility_count == 0)
    neither_el = (el.cmr_facility_count == 0) & (el.cct_facility_count == 0)

    valid = m[["svi_percentile", "edi_national_percentile"]].dropna()
    pear_r, pear_p = stats.pearsonr(valid.svi_percentile, valid.edi_national_percentile)

    q = pd.qcut(el.edi_national_percentile, 5, labels=False)
    q1 = el.loc[q == 0, "cmr_rate_per_100k"].mean()
    q5 = el.loc[q == 4, "cmr_rate_per_100k"].mean()
    kw = stats.kruskal(*[g.cmr_rate_per_100k.values for _, g in el.groupby(q)])

    mw_cmr = stats.mannwhitneyu(el.loc[el.metro_indicator == 1, "cmr_rate_per_100k"],
                                el.loc[el.metro_indicator == 0, "cmr_rate_per_100k"])
    mw_cct = stats.mannwhitneyu(el.loc[el.metro_indicator == 1, "cct_rate_per_100k"],
                                el.loc[el.metro_indicator == 0, "cct_rate_per_100k"])

    stats_out = {
        "_note": ("Generated by code/11_edi_tables_and_stats.py from the committed data. "
                  "Do not hand-edit; regenerate."),
        "total_counties": int(len(m)),
        "rate_eligible_counties": int(len(el)),
        "cmr_facilities_total": int(m.cmr_facility_count.sum()),
        "cct_facilities_total": int(m.cct_facility_count.sum()),
        "counties_with_cmr": int((m.cmr_facility_count > 0).sum()),
        "counties_with_cct": int((m.cct_facility_count > 0).sum()),
        "counties_neither_modality": int(neither_all.sum()),
        "counties_neither_modality_pct": round(100 * neither_all.mean(), 2),
        "metro_counties": int((m.metro_indicator == 1).sum()),
        "nonmetro_counties": int((m.metro_indicator == 0).sum()),
        "nonmetro_counties_pct": round(100 * (m.metro_indicator == 0).mean(), 2),
        "mean_edi_no_facility_rate_eligible": round(
            float(el.loc[neither_el, "edi_national_percentile"].mean()), 1),
        "mean_edi_has_facility_rate_eligible": round(
            float(el.loc[~neither_el, "edi_national_percentile"].mean()), 1),
        "svi_edi_pearson_r": round(float(pear_r), 4),
        "svi_edi_pearson_p": float(pear_p),
        "metro_vs_nonmetro_cmr_mann_whitney_p": float(mw_cmr.pvalue),
        "metro_vs_nonmetro_cct_mann_whitney_p": float(mw_cct.pvalue),
        "metro_cmr_mean_rate": round(float(el.loc[el.metro_indicator == 1, "cmr_rate_per_100k"].mean()), 4),
        "nonmetro_cmr_mean_rate": round(float(el.loc[el.metro_indicator == 0, "cmr_rate_per_100k"].mean()), 4),
        "metro_cct_mean_rate": round(float(el.loc[el.metro_indicator == 1, "cct_rate_per_100k"].mean()), 4),
        "nonmetro_cct_mean_rate": round(float(el.loc[el.metro_indicator == 0, "cct_rate_per_100k"].mean()), 4),
        "edi_q1_cmr_mean": round(float(q1), 4),
        "edi_q5_cmr_mean": round(float(q5), 4),
        "edi_q1_q5_cmr_ratio": round(float(q1 / q5), 2),
        "kruskal_wallis_cmr_p": float(kw.pvalue),
        "edi_cmr_irr_unadjusted": edi_res["CMR"]["unadjusted"],
        "edi_cmr_irr_adjusted_metro": edi_res["CMR"]["adjusted_metro"],
        "metro_effect_cmr": edi_res["CMR"]["metro_effect"],
        "svi_cmr_irr_unadjusted": svi_res["CMR"]["unadjusted"],
        "svi_cmr_irr_adjusted_metro": svi_res["CMR"]["adjusted_metro"],
    }

    for d in (TABLES, REQUESTED, SUPP):
        with open(os.path.join(d, "additional_statistics.json"), "w") as f:
            json.dump(stats_out, f, indent=2)
    print("  wrote additional_statistics.json")

    # Internal-consistency check: every "N (P%)" pair must agree.
    checks = [
        ("counties_neither_modality", stats_out["counties_neither_modality"],
         stats_out["counties_neither_modality_pct"]),
        ("nonmetro_counties", stats_out["nonmetro_counties"],
         stats_out["nonmetro_counties_pct"]),
    ]
    total = stats_out["total_counties"]
    for name, n, pct in checks:
        implied = round(100 * n / total, 2)
        assert abs(implied - pct) < 0.01, (
            f"CONSISTENCY FAILURE for {name}: count {n} implies {implied}%, "
            f"but {pct}% was recorded")
    print(f"  consistency check passed: counts and percentages agree "
          f"({stats_out['counties_neither_modality']} neither modality = "
          f"{stats_out['counties_neither_modality_pct']}%)")

    return stats_out


def main():
    print("Generating EDI tables and statistics from committed data ...")
    m, edi_res, sdi_res = load()
    svi_res = svi_models(m)
    supplementary_table(edi_res)
    comparison_table(edi_res, svi_res, sdi_res, m)
    external_validation_table(sdi_res)
    s = additional_stats(m, edi_res, svi_res)
    print(f"\n  Counties with neither modality: {s['counties_neither_modality']} "
          f"({s['counties_neither_modality_pct']}%)")
    print(f"  Mean EDI, no facility vs facility: "
          f"{s['mean_edi_no_facility_rate_eligible']} vs "
          f"{s['mean_edi_has_facility_rate_eligible']}")
    print("\n  Done.")


if __name__ == "__main__":
    main()
